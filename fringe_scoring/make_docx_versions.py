"""简明说明 / 测试集检测报告 的 Word（.docx）版本生成器（开发壳，不随核心包交付）。

复用两份 PDF 生成器的页面块结构（brief_pages / report_pages，图文同源），
映射为 Word 构件：标题/正文/注释为可编辑文本，含 mathtext 的公式行渲染为
透明 PNG 嵌入（与 PDF 同一排版引擎，视觉一致），"式中"注释里的行内数学
符号转为 Unicode 文本，全角空格分隔的表格转为真正的 Word 表格。
用法：venv python fringe_scoring/make_docx_versions.py
"""

from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # 脚本直跑时可 import 包

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from fringe_scoring.make_batch_report_doc import report_pages
from fringe_scoring.make_position_brief_doc import brief_pages

ROOT = Path(__file__).resolve().parents[1]

try:  # Windows 控制台默认 GBK，强制 UTF-8
    sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[union-attr]
except Exception:
    pass

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei"]
plt.rcParams["axes.unicode_minus"] = False

# 行内 mathtext → Unicode 文本（"式中"注释/表格用；公式本体整行转图片不走此表）
_TEX_MAP = [
    (r"\\rho_u", "ρu"), (r"\\rho_0", "ρ0"), (r"\\sigma_\{?ref\}?", "σ_ref"),
    (r"\\sigma_R", "σ_R"), (r"\\Omega_u", "Ω_u"), (r"\\Omega", "Ω"),
    (r"\\mathrm\{([^}]*)\}", r"\1"), (r"\\max", "max"), (r"\\min", "min"),
    (r"\\geq", "≥"), (r"\\leq", "≤"), (r"\\in\b", "∈"), (r"\\times", "×"),
    (r"\\equiv", "≡"), (r"\\cdot", "·"), (r"U_\\rho", "U_ρ"),
    (r"Q_\{([^}]*)\}", r"Q_\1"), (r"_\{([^}]*)\}", r"_\1"),
    (r"\\,|\\;|\\ ", " "), (r"[{}]", ""),
]


def tex_to_text(s: str) -> str:
    """去掉 $ 定界并把行内 mathtext 命令换成 Unicode 近似文本。"""
    out = s.replace("$", "")
    for pat, rep in _TEX_MAP:
        out = re.sub(pat, rep, out)
    return out


def _set_font(run, size_pt: float, bold: bool = False, gray: bool = False) -> None:
    """run 字体统一：微软雅黑（含中文 eastAsia 字体项）。"""
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size_pt)
    run.bold = bold
    if gray:
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _para(doc, text: str, size: float, bold=False, gray=False, center=False, indent=0.0):
    """加一段（多行文本按行拆 run + 手工换行），返回段落。"""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        _set_font(p.add_run(line), size, bold=bold, gray=gray)
    return p


def _formula_png(text: str, tmpdir: Path, idx: int) -> Path:
    """公式行（中文前缀 + mathtext）→ 紧致透明 PNG，与 PDF 同渲染引擎。"""
    fig = plt.figure(figsize=(8, 1))
    fig.text(0.5, 0.5, text, fontsize=13, ha="center", va="center")
    out = tmpdir / f"formula_{idx:03d}.png"
    fig.savefig(out, dpi=220, bbox_inches="tight", pad_inches=0.06, transparent=True)
    plt.close(fig)
    return out


def _add_centered_picture(doc, path: Path, width_in: float) -> None:
    """居中插图。"""
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _png_width_in(path: Path, dpi: float, cap: float) -> float:
    """按 PNG 像素尺寸换算插入宽度（英寸），封顶 cap。"""
    from PIL import Image

    with Image.open(path) as im:
        return min(cap, im.width / dpi)


def _add_table(doc, tbl_text: str) -> None:
    """全角空格分隔的对齐文本表 → Word 表格（首行加粗，网格线）。"""
    rows = [[c for c in re.split(r"　+", line.strip("　")) if c]
            for line in tbl_text.split("\n")]
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(tex_to_text(row[ci]) if ci < len(row) else "")
            _set_font(run, 9, bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_imgrow(doc, spec: str) -> None:
    """双图并排 + 题注：借一张无边框 1×2 表格排版。"""
    p1, p2, _h, c1, c2 = spec.split("::")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (img, cap) in enumerate(((p1, c1), (p2, c2))):
        cell = table.cell(0, ci)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(ROOT / img), width=Inches(2.9))
        cap_p = table.cell(1, ci).paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(cap_p.add_run(cap), 9, gray=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_docx(pages: list[list[tuple[str, str]]], out_path: Path) -> None:
    """页面块序列 → .docx（Word 自动分栏排版，忽略 PDF 的手工分页与留白）。"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)
    with tempfile.TemporaryDirectory() as td:
        tmpdir = Path(td)
        fi = 0
        for blocks in pages:
            for kind, text in blocks:
                if kind == "space":
                    continue
                if kind == "h1":
                    _para(doc, text, 16, bold=True, center=True)
                elif kind == "sub":
                    _para(doc, text, 12, center=True)
                elif kind == "cnote":
                    _para(doc, text, 9, gray=True, center=True)
                elif kind == "h2":
                    p = _para(doc, text, 13, bold=True)
                    p.paragraph_format.space_before = Pt(10)
                elif kind == "h3":
                    p = _para(doc, text, 11, bold=True)
                    p.paragraph_format.space_before = Pt(6)
                elif kind == "formula":
                    fi += 1
                    png = _formula_png(text, tmpdir, fi)
                    _add_centered_picture(doc, png, _png_width_in(png, 220, 6.2))
                elif kind == "vars":
                    _para(doc, "式中：", 9.5, bold=True)
                    _para(doc, tex_to_text(text), 9.5, indent=0.35)
                elif kind == "tbl":
                    _add_table(doc, text)
                elif kind == "note":
                    _para(doc, text, 9, gray=True)
                elif kind == "img":
                    rel, _h = text.rsplit("::", 1)
                    _add_centered_picture(doc, ROOT / rel, 5.6)
                elif kind == "imgrow":
                    _add_imgrow(doc, text)
                else:  # body 及其余文本块
                    _para(doc, tex_to_text(text), 10.5)
    doc.save(str(out_path))
    print(f"已生成 → {out_path}")


def main() -> int:
    """生成两份 Word 版到 docs/。"""
    build_docx(brief_pages(), ROOT / "docs" / "应力斑位置评分指标_简明说明.docx")
    build_docx(report_pages(), ROOT / "docs" / "应力斑位置评分_测试集检测报告.docx")
    return 0


if __name__ == "__main__":
    sys.exit(main())
