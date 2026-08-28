"""专利交底书 Markdown → 正式 docx 转换器（开发壳，不随核心包交付）。

输入：命令行第 1 参数指定 md 路径；缺省为 docs/专利交底书_钢化炉逆向诊断_正式版.md
（内容单一事实源，改 md 后重跑即可）。输出：同目录同名 .docx。

排版口径（正式、无花哨元素）：A4；正文宋体小四 1.5 倍行距、首行缩进 2 字符；
标题黑体；引用块楷体；表格 Table Grid 五号；公式行居中；附图居中宽 14.5cm。
解析的 Markdown 子集：#/##/### 标题、表格、图片、引用块、列表行、
以"  "开头且含"……("的公式行、**加粗** 内联、--- 分隔线（忽略）。
用法：venv python fringe_scoring/make_patent_docx.py [md路径]
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

MD_PATH = (Path(sys.argv[1]).resolve() if len(sys.argv) > 1
           else ROOT / "docs" / "专利交底书_钢化炉逆向诊断_正式版.md")
OUT_PATH = MD_PATH.with_suffix(".docx")

try:  # Windows 控制台默认 GBK，强制 UTF-8
    sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[union-attr]
except Exception:
    pass

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")


def _set_run(run, cn_font: str, size_pt: float, bold: bool = False) -> None:
    """设置 run 的中英文字体（中文按 eastAsia，西文 Times New Roman）、字号、加粗。"""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def _add_runs(par, text: str, cn_font: str, size_pt: float, base_bold: bool = False) -> None:
    """按 **加粗** 内联标记把一行文本拆成多个 run 写入段落。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _set_run(par.add_run(text[pos:m.start()]), cn_font, size_pt, base_bold)
        _set_run(par.add_run(m.group(1)), cn_font, size_pt, True)
        pos = m.end()
    if pos < len(text):
        _set_run(par.add_run(text[pos:]), cn_font, size_pt, base_bold)


def _body_par(doc, text: str, indent: bool = True, cn_font: str = "宋体",
              size_pt: float = 12.0, align=None, bold: bool = False):
    """添加一个正文段落（默认宋体小四、1.5 倍行距、首行缩进 2 字符）。"""
    par = doc.add_paragraph()
    fmt = par.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        fmt.first_line_indent = Pt(size_pt * 2)  # 首行缩进 2 字符
    if align is not None:
        fmt.alignment = align
    _add_runs(par, text, cn_font, size_pt, bold)
    return par


def _heading(doc, text: str, level: int) -> None:
    """标题：level 0=文档主标题（三号黑体居中），1=## 四号黑体，2=### 小四黑体。"""
    par = doc.add_paragraph()
    fmt = par.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if level == 0:
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(12)
        size = 16.0
    else:
        fmt.space_before = Pt(10 if level == 1 else 6)
        fmt.space_after = Pt(4)
        size = 14.0 if level == 1 else 12.0
    _add_runs(par, text, "黑体", size, base_bold=True)


def _split_row(line: str) -> list[str]:
    """把一行 Markdown 表格拆成单元格文本（支持 \\| 转义）。"""
    cells = re.split(r"(?<!\\)\|", line.strip().strip("|"))
    return [c.replace("\\|", "|").strip() for c in cells]


def _add_table(doc, rows: list[list[str]]) -> None:
    """Markdown 表格 → 带框线 docx 表格（表头加粗，五号字，单倍行距）。"""
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            text = row[ci] if ci < len(row) else ""
            _add_runs(par, text, "宋体", 10.5, base_bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)  # 表后空行，防表格粘连


def _add_image(doc, alt: str, rel_path: str) -> bool:
    """居中插入附图（宽 14.5cm）+ 图题段；路径相对 md 所在目录解析。"""
    img = (MD_PATH.parent / rel_path).resolve()
    if not img.exists():
        print(f"  ⚠️ 附图缺失，跳过：{rel_path}")
        return False
    par = doc.add_paragraph()
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(img), width=Cm(14.5))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    _set_run(cap.add_run(alt), "宋体", 10.5)
    return True


def _setup_page(doc) -> None:
    """A4 页面与页边距（上下 2.54cm、左右 3.0cm）。"""
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.0)


def convert() -> tuple[int, int, int]:
    """md → docx 主流程；返回（段落数, 表格数, 附图数）供自检。"""
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    _setup_page(doc)

    lines = md.split("\n")
    i, n_img = 0, 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            _heading(doc, line[2:].strip(), 0)
        elif line.startswith("## "):
            _heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            _heading(doc, line[4:].strip(), 2)
        elif line.lstrip().startswith(">"):
            # 卷首说明引用块：楷体五号、左缩进，不首行缩进
            text = line.lstrip()[1:].strip()
            if text:
                par = _body_par(doc, text, indent=False, cn_font="楷体", size_pt=10.5)
                par.paragraph_format.left_indent = Cm(0.75)
        elif _IMG_RE.match(line.strip()):
            m = _IMG_RE.match(line.strip())
            assert m is not None
            if _add_image(doc, m.group(1), m.group(2)):
                n_img += 1
        elif line.strip().startswith("|"):
            # 表格块：收集连续 | 行，跳过对齐分隔行
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = _split_row(lines[i])
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            _add_table(doc, rows)
            continue
        elif line.startswith("- "):
            par = _body_par(doc, line[2:].strip(), indent=False)
            par.paragraph_format.left_indent = Cm(0.75)
        elif line.startswith("  ") and "……(" in line:
            _body_par(doc, line.strip(), indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _body_par(doc, line.strip())
        i += 1

    doc.save(OUT_PATH)
    return len(doc.paragraphs), len(doc.tables), n_img


def verify() -> None:
    """读回自检：段落/表格/内嵌图数量与文件大小（python-docx 能重新打开即结构完整）。"""
    doc = Document(str(OUT_PATH))
    n_shapes = len(doc.inline_shapes)
    size_kb = OUT_PATH.stat().st_size / 1024
    print(f"  读回自检：段落 {len(doc.paragraphs)}，表格 {len(doc.tables)}，"
          f"内嵌图 {n_shapes}，文件 {size_kb:.0f} KB")


def main() -> None:
    """转换并自检。"""
    if not MD_PATH.exists():
        raise FileNotFoundError(f"未找到 Markdown 源：{MD_PATH}")
    n_par, n_tab, n_img = convert()
    print(f"已生成：{OUT_PATH.relative_to(ROOT)}（段落 {n_par}，表格 {n_tab}，附图 {n_img}）")
    verify()


if __name__ == "__main__":
    main()
